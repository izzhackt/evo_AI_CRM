#!/usr/bin/env python3
"""Build the branded EVO Platform technical specification DOCX.

Design contract:
  preset: standard_business_brief
  page: US Letter, portrait, 1-inch margins, 0.492-inch header/footer
  body: Arial 11pt (named brand/runtime override from Calibri), 1.10 spacing,
        6pt after
  headings: preset sizes/spacing with EVO brand color overrides
  lists: marker 0.25in, text 0.5in, hanging 0.25in, 8pt after, 1.167 spacing
  tables: 9360 DXA, indent 120 DXA, fixed geometry, 80/120 cell margins

Named visual overrides:
  - Arial replaces Calibri because the approved brand is sans-serif and Arial
    is reliably available in the document render/runtime.
  - EVO red #D70217 replaces preset blue for Heading 1 and restrained accents.
  - Heading 2/3 use black/gray to avoid a dense "wall of red".
  - The first page follows the memo_masthead pattern with the approved logo.
"""

from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Iterable, Sequence
from urllib.parse import urlparse
from zipfile import ZipFile, ZipInfo

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap, qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/specs/EVO_PLATFORM_TZ.md"
OUTPUT = ROOT / "docs/specs/EVO_PLATFORM_TZ.docx"
LOGOBOOK = ROOT / "docs/company/brand/evo-admissions-logobook.pdf"

PRESET_NAME = "standard_business_brief"
FONT = "Arial"
MONO_FONT = "Courier New"
PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
MARGIN_IN = 1.0
CONTENT_WIDTH_IN = 6.5
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_START_END_DXA = 120

EVO_RED = "D70217"
EVO_RED_DARK = "A60017"
INK = "111111"
GRAY = "666666"
MUTED = "72777F"
LIGHT_GRAY = "F2F4F7"
LIGHT_RED = "FBEAEC"
DOCX_ZIP_TIMESTAMP = (2026, 7, 26, 0, 0, 0)
WHITE = "FFFFFF"
BORDER = "D9DDE3"
CODE_FILL = "F6F7F9"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(
    run,
    *,
    name: str = FONT,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)
    r_fonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(
    paragraph,
    *,
    before: float = 0,
    after: float = 6,
    line: float = 1.10,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_left_border(
    paragraph,
    *,
    color: str = EVO_RED,
    size: str = "18",
    space: str = "8",
) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)


def paragraph_bottom_border(paragraph, *, color: str = EVO_RED, size: str = "20") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (
        ("top", CELL_TOP_BOTTOM_DXA),
        ("bottom", CELL_TOP_BOTTOM_DXA),
        ("start", CELL_START_END_DXA),
        ("end", CELL_START_END_DXA),
    ):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def widths_for(headers: Sequence[str]) -> list[int]:
    n = len(headers)
    lowered = [h.lower() for h in headers]
    if n == 2:
        return [2700, 6660]
    if n == 3:
        return [1350, 5100, 2910]
    if n == 4:
        if "цель / правило" in lowered:
            return [1150, 1800, 4000, 2410]
        if "кто утверждает" in lowered:
            return [1150, 4900, 1900, 1410]
        if lowered[0] == "id":
            return [1150, 4650, 1200, 2360]
        if "что хранит evo platform" in lowered:
            return [1800, 3000, 2200, 2360]
        return [1800, 3480, 1440, 2640]
    if n == 5:
        return [1800, 1800, 2400, 1080, 2280]
    if n == 6:
        return [1800, 1512, 1512, 1512, 1512, 1512]
    if n == 7:
        return [1800, 1260, 1260, 1260, 1260, 1260, 1260]
    base = CONTENT_WIDTH_DXA // n
    widths = [base] * n
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def hyperlink(paragraph, text: str, url: str, *, color: str = EVO_RED_DARK) -> None:
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    r_color = OxmlElement("w:color")
    r_color.set(qn("w:val"), color)
    r_pr.append(r_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT)
    r_pr.append(r_fonts)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    link.append(run)
    paragraph._p.append(link)


INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|https?://[^\s]+)"
)


def add_inline(paragraph, text: str, *, size: float = 11, color: str = INK) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name=MONO_FONT, size=max(size - 0.5, 7.5), color=INK)
            run.font.highlight_color = None
        elif token.startswith("["):
            link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_match:
                label, url = link_match.groups()
                if url.startswith(("http://", "https://")):
                    hyperlink(paragraph, label, url)
                else:
                    run = paragraph.add_run(label)
                    set_run_font(run, size=size, color=EVO_RED_DARK)
            else:
                run = paragraph.add_run(token)
                set_run_font(run, size=size, color=color)
        else:
            trimmed = token.rstrip(".,);")
            suffix = token[len(trimmed) :]
            hyperlink(paragraph, trimmed, trimmed)
            if suffix:
                run = paragraph.add_run(suffix)
                set_run_font(run, size=size, color=color)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color)


def add_page_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_section(section) -> None:
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_header_footer(section) -> None:
    section.different_first_page_header_footer = True
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, after=0, line=1.0)
    left = p.add_run("EVO Admissions  |  Техническое задание")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    p.add_run("\t")
    right = p.add_run("Версия 1.0  ·  26.07.2026")
    set_run_font(right, size=8.5, color=MUTED)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(CONTENT_WIDTH_IN), WD_ALIGN_PARAGRAPH.RIGHT)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(fp, after=0, line=1.0)
    left_f = fp.add_run("Внутренний документ EVO Admissions")
    set_run_font(left_f, size=8.5, color=MUTED)
    fp.add_run("\t")
    page_label = fp.add_run("Стр. ")
    set_run_font(page_label, size=8.5, color=MUTED)
    add_page_field(fp, "PAGE")
    of_run = fp.add_run(" из ")
    set_run_font(of_run, size=8.5, color=MUTED)
    add_page_field(fp, "NUMPAGES")
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(CONTENT_WIDTH_IN), WD_ALIGN_PARAGRAPH.RIGHT)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    h1 = styles["Heading 1"]
    h1.font.name = FONT
    h1._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = rgb(EVO_RED)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = FONT
    h2._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = rgb(INK)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.0
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = FONT
    h3._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = rgb(GRAY)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.line_spacing = 1.0
    h3.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = FONT
    caption._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def extract_logo(tmp_dir: Path) -> Path:
    pdftoppm = shutil.which("pdftoppm")
    if not pdftoppm:
        raise RuntimeError("pdftoppm is required to extract the approved logo")
    prefix = tmp_dir / "evo-logo"
    subprocess.run(
        [
            pdftoppm,
            "-png",
            "-singlefile",
            "-r",
            "120",
            "-f",
            "5",
            "-l",
            "5",
            "-x",
            "390",
            "-y",
            "390",
            "-W",
            "1020",
            "-H",
            "660",
            str(LOGOBOOK),
            str(prefix),
        ],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.PIPE,
    )
    path = prefix.with_suffix(".png")
    if not path.exists():
        raise RuntimeError("Approved logo extraction did not create a PNG")
    return path


def add_cover(document: Document, logo_path: Path) -> None:
    for _ in range(2):
        p = document.add_paragraph()
        set_paragraph_spacing(p, after=0)

    logo_p = document.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_run = logo_p.add_run()
    shape = logo_run.add_picture(str(logo_path), width=Inches(3.5))
    shape._inline.docPr.set("descr", "Официальный логотип EVO Admissions")
    set_paragraph_spacing(logo_p, after=24)

    kicker = document.add_paragraph()
    set_paragraph_spacing(kicker, after=8, line=1.0)
    run = kicker.add_run("ТЕХНИЧЕСКОЕ ЗАДАНИЕ")
    set_run_font(run, size=11, color=EVO_RED, bold=True)

    title = document.add_paragraph()
    set_paragraph_spacing(title, after=8, line=1.0)
    run = title.add_run("Единая платформа\nавтоматизации EVO Admissions")
    set_run_font(run, size=25, color=INK, bold=True)

    subtitle = document.add_paragraph()
    set_paragraph_spacing(subtitle, after=18, line=1.15)
    run = subtitle.add_run(
        "Продажи · WhatsApp · Поступление · Документы · Виза · Финансы · Student Portal"
    )
    set_run_font(run, size=12.5, color=GRAY)

    rule = document.add_paragraph()
    set_paragraph_spacing(rule, after=18, line=1.0)
    paragraph_bottom_border(rule, color=EVO_RED, size="28")

    metadata = [
        ("Документ", "EVO-PLATFORM-TZ-001"),
        ("Версия", "1.0"),
        ("Статус", "Проект для согласования владельцем бизнеса"),
        ("Дата", "26 июля 2026 года"),
        ("Базовый commit", "0ecd95d6b248572269bec17d60072a49230e626e"),
    ]
    for label, value in metadata:
        p = document.add_paragraph()
        set_paragraph_spacing(p, after=3, line=1.0)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=10.5, color=INK, bold=True)
        value_run = p.add_run(value)
        set_run_font(
            value_run,
            name=MONO_FONT if label == "Базовый commit" else FONT,
            size=9 if label == "Базовый commit" else 10.5,
            color=INK,
        )

    document.add_paragraph()
    notice = document.add_paragraph()
    set_paragraph_spacing(notice, before=8, after=0, line=1.15)
    paragraph_shading(notice, LIGHT_RED)
    paragraph_left_border(notice)
    add_inline(
        notice,
        "Это ТЗ является контрактом на последующую реализацию. "
        "Backend-миграция начинается только после утверждения открытых решений.",
        size=10.5,
    )
    document.add_page_break()


def add_table_of_contents(document: Document) -> None:
    toc_h = document.add_paragraph("Содержание", style="Heading 1")
    toc_h.paragraph_format.space_before = Pt(0)
    toc = document.add_paragraph()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Оглавление обновляется при открытии документа."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = toc.add_run()
    run._r.extend([begin, instr, separate, placeholder, end])
    set_run_font(run, size=10.5, color=MUTED)
    document.add_page_break()


def create_numbering(document: Document, *, ordered: bool) -> int:
    numbering = document.part.numbering_part.element
    existing_abstract = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    abstract_id = max(existing_abstract, default=-1) + 1
    existing_num = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    lvl.append(lvl_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    lvl.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    if not ordered:
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), FONT)
        fonts.set(qn("w:hAnsi"), FONT)
        r_pr.append(fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(document: Document, text: str, num_id: int) -> None:
    p = document.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    add_inline(p, text)


def add_table(document: Document, rows: Sequence[Sequence[str]]) -> None:
    if not rows:
        return
    headers = list(rows[0])
    widths = widths_for(headers)
    table = document.add_table(rows=len(rows), cols=len(headers))
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    font_size = 9.2 if len(headers) <= 3 else 8.5
    if len(headers) >= 6:
        font_size = 7.6
    elif len(headers) == 5:
        font_size = 8.0

    for row_idx, values in enumerate(rows):
        row = table.rows[row_idx]
        for col_idx, value in enumerate(values):
            cell = row.cells[col_idx]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            set_paragraph_spacing(paragraph, after=0, line=1.05)
            if row_idx == 0:
                shade_cell(cell, LIGHT_GRAY)
            add_inline(paragraph, value, size=font_size, color=INK)
            for run in paragraph.runs:
                if row_idx == 0:
                    run.bold = True
    spacer = document.add_paragraph()
    set_paragraph_spacing(spacer, after=2, line=1.0)


def add_callout(document: Document, text: str) -> None:
    p = document.add_paragraph()
    set_paragraph_spacing(p, before=4, after=8, line=1.15)
    paragraph_shading(p, LIGHT_RED)
    paragraph_left_border(p)
    p.paragraph_format.left_indent = Inches(0.10)
    p.paragraph_format.right_indent = Inches(0.08)
    add_inline(p, text, size=10.5)


def add_code_block(document: Document, lines: Sequence[str]) -> None:
    p = document.add_paragraph()
    set_paragraph_spacing(p, before=4, after=8, line=1.0)
    p.paragraph_format.keep_together = True
    paragraph_shading(p, CODE_FILL)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, name=MONO_FONT, size=8.3, color=INK)
        if idx != len(lines) - 1:
            run.add_break()


MAJOR_PAGE_BREAK_PREFIXES = (
    "6.",
    "13.",
    "27.",
    "32.",
)


def add_heading(document: Document, text: str, level: int) -> None:
    style = f"Heading {level}"
    p = document.add_paragraph(style=style)
    if level == 1 and text.startswith(MAJOR_PAGE_BREAK_PREFIXES):
        p.paragraph_format.page_break_before = True
    if level == 2 and text.startswith("27.") and not text.startswith("27.1 "):
        p.paragraph_format.page_break_before = True
    add_inline(
        p,
        text,
        size=16 if level == 1 else 13 if level == 2 else 12,
        color=EVO_RED if level == 1 else INK if level == 2 else GRAY,
    )
    for run in p.runs:
        run.bold = True


def add_image(document: Document, alt: str, path_text: str) -> None:
    path = (SOURCE.parent / path_text).resolve()
    if not path.exists():
        raise FileNotFoundError(f"Screenshot not found: {path}")
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=2, line=1.0)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(6.25))
    shape._inline.docPr.set("descr", alt)
    caption = document.add_paragraph(style="Caption")
    add_inline(caption, alt, size=9, color=MUTED)


SPECIAL_RE = re.compile(
    r"^(#{1,6}\s|[-*]\s|\d+\.\s|>\s?|```|\|.*\||!\[)"
)


def parse_table(lines: Sequence[str], start: int) -> tuple[list[list[str]], int]:
    raw: list[str] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw.append(lines[i].strip())
        i += 1
    parsed: list[list[str]] = []
    for idx, line in enumerate(raw):
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if idx == 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        parsed.append(cells)
    expected = len(parsed[0]) if parsed else 0
    for row in parsed:
        if len(row) != expected:
            raise ValueError(f"Malformed Markdown table near line {start + 1}")
    return parsed, i


def parse_markdown(document: Document) -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("## 1. "))
    i = start
    active_list_kind: str | None = None
    active_num_id: int | None = None

    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            active_list_kind = None
            active_num_id = None
            i += 1
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            marks, text = heading.groups()
            add_heading(document, text, len(marks) - 1)
            active_list_kind = None
            active_num_id = None
            i += 1
            continue

        if stripped.startswith("```"):
            code: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i].rstrip())
                i += 1
            if i < len(lines):
                i += 1
            add_code_block(document, code)
            active_list_kind = None
            active_num_id = None
            continue

        if stripped.startswith(">"):
            parts: list[str] = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                parts.append(lines[i].strip().lstrip(">").strip())
                i += 1
            add_callout(document, " ".join(parts))
            active_list_kind = None
            active_num_id = None
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(document, rows)
            active_list_kind = None
            active_num_id = None
            continue

        image = re.match(r"^!\[([^\]]+)\]\(([^)]+)\)$", stripped)
        if image:
            alt, path_text = image.groups()
            add_image(document, alt, path_text)
            i += 1
            active_list_kind = None
            active_num_id = None
            continue

        list_match = re.match(r"^([-*]|\d+\.)\s+(.+)$", stripped)
        if list_match:
            marker, text = list_match.groups()
            kind = "ordered" if marker[0].isdigit() else "bullet"
            i += 1
            continuations: list[str] = []
            while i < len(lines):
                candidate = lines[i].strip()
                if not candidate or SPECIAL_RE.match(candidate):
                    break
                continuations.append(candidate)
                i += 1
            if continuations:
                text = " ".join([text, *continuations])
            if active_list_kind != kind or active_num_id is None:
                active_num_id = create_numbering(document, ordered=kind == "ordered")
                active_list_kind = kind
            add_list_item(document, text, active_num_id)
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            candidate = lines[i].strip()
            if not candidate or SPECIAL_RE.match(candidate):
                break
            paragraph_lines.append(candidate)
            i += 1
        p = document.add_paragraph()
        set_paragraph_spacing(p, after=6, line=1.10)
        add_inline(p, " ".join(paragraph_lines))
        active_list_kind = None
        active_num_id = None


def set_document_properties(document: Document) -> None:
    props = document.core_properties
    props.title = "Техническое задание — Единая платформа EVO Admissions"
    props.subject = "Архитектура, функции, доступ, интеграции, миграция и приёмка"
    props.author = "EVO Admissions"
    props.keywords = "EVO Admissions, amoCRM, Supabase, WAHA, WhatsApp, ТЗ"
    props.comments = "Generated from docs/specs/EVO_PLATFORM_TZ.md"


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def audit_document(document: Document) -> None:
    for section in document.sections:
        assert section.page_width == Inches(PAGE_WIDTH_IN)
        assert section.page_height == Inches(PAGE_HEIGHT_IN)
        assert section.left_margin == Inches(MARGIN_IN)
        assert section.right_margin == Inches(MARGIN_IN)
        assert section.top_margin == Inches(MARGIN_IN)
        assert section.bottom_margin == Inches(MARGIN_IN)

    normal = document.styles["Normal"]
    assert normal.font.name == FONT
    assert normal.font.size == Pt(11)
    assert normal.paragraph_format.space_after == Pt(6)
    assert abs(float(normal.paragraph_format.line_spacing) - 1.10) < 0.001

    for table in document.tables:
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        assert tbl_w is not None and tbl_w.get(qn("w:w")) == str(CONTENT_WIDTH_DXA)
        assert tbl_w.get(qn("w:type")) == "dxa"
        assert tbl_ind is not None and tbl_ind.get(qn("w:w")) == str(TABLE_INDENT_DXA)
        grid_widths = [
            int(col.get(qn("w:w"))) for col in table._tbl.tblGrid.findall(qn("w:gridCol"))
        ]
        assert sum(grid_widths) == CONTENT_WIDTH_DXA
        assert table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader")) is not None


def normalize_docx_package(path: Path) -> None:
    """Rewrite the DOCX ZIP with fixed metadata for bit-for-bit builds."""
    normalized = path.with_name(f".{path.name}.normalized")
    try:
        with ZipFile(path, "r") as source, ZipFile(normalized, "w") as target:
            for member in source.infolist():
                info = ZipInfo(member.filename, date_time=DOCX_ZIP_TIMESTAMP)
                info.compress_type = member.compress_type
                info.comment = member.comment
                info.create_system = member.create_system
                info.external_attr = member.external_attr
                info.internal_attr = member.internal_attr
                target.writestr(
                    info,
                    source.read(member.filename),
                    compress_type=member.compress_type,
                    compresslevel=9,
                )
        normalized.replace(path)
    finally:
        normalized.unlink(missing_ok=True)


def build() -> Path:
    document = Document()
    for section in document.sections:
        configure_section(section)
        configure_header_footer(section)
    configure_styles(document)
    set_document_properties(document)
    set_update_fields(document)

    with tempfile.TemporaryDirectory(prefix="evo-tz-docx-") as tmp:
        logo = extract_logo(Path(tmp))
        add_cover(document, logo)
        add_table_of_contents(document)
        parse_markdown(document)

    audit_document(document)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    normalize_docx_package(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build()
    print(output)
