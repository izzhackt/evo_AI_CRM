#!/usr/bin/env python3
"""Build and verify the owner-facing EVO Platform technical specification.

The verifier uses only real document paths and real renderers:

- python-docx builds and inspects the DOCX package;
- LibreOffice/soffice converts the DOCX to PDF;
- Poppler/pdftoppm renders every PDF page to PNG.

It writes machine-readable validation and accessibility reports beside the
rendered pages. Human visual inspection of every generated PNG remains a
separate, explicit release gate and is recorded in the validation ledger.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

import docx
from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/specs/EVO_PLATFORM_TZ.md"
DOCX = ROOT / "docs/specs/EVO_PLATFORM_TZ.docx"
GENERATOR = ROOT / "scripts/generate-evo-platform-tz.py"

EXPECTED_REQUIREMENTS = {
    "FR": 90,
    "INT": 20,
    "DATA": 18,
    "SEC": 20,
    "NFR": 18,
    "ACC": 25,
    "DEC": 20,
}

REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def run(command: list[str]) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        command,
        check=True,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )


def require_command(*names: str) -> str:
    for name in names:
        resolved = shutil.which(name)
        if resolved:
            return resolved
    joined = " or ".join(names)
    raise RuntimeError(f"Required real renderer command is missing: {joined}")


def expected_requirement_ids() -> set[str]:
    return {
        f"{prefix}-{number:03d}"
        for prefix, maximum in EXPECTED_REQUIREMENTS.items()
        for number in range(1, maximum + 1)
    }


def validate_markdown() -> dict[str, Any]:
    text = SOURCE.read_text(encoding="utf-8")
    expected = expected_requirement_ids()

    all_ids = re.findall(r"\b(FR|INT|DATA|SEC|NFR|ACC|DEC)-(\d{3})\b", text)
    unique_ids = {f"{prefix}-{number}" for prefix, number in all_ids}
    missing_from_document = sorted(expected - unique_ids)
    unexpected_ids = sorted(unique_ids - expected)
    if missing_from_document or unexpected_ids:
        raise RuntimeError(
            "Requirement ID coverage mismatch: "
            f"missing={missing_from_document}, unexpected={unexpected_ids}"
        )

    try:
        traceability = text.split(
            "### 31.4 Поэлементная матрица происхождения требований", 1
        )[1].split("## 32. Согласование", 1)[0]
    except IndexError as exc:
        raise RuntimeError("Traceability section 31.4 is missing or malformed") from exc

    trace_rows = re.findall(
        r"^\| ((?:FR|INT|DATA|SEC|NFR|ACC|DEC)-\d{3}) \| ([^|]+) \|$",
        traceability,
        flags=re.MULTILINE,
    )
    trace_counts = Counter(requirement_id for requirement_id, _ in trace_rows)
    duplicates = sorted(key for key, count in trace_counts.items() if count != 1)
    missing_trace = sorted(expected - set(trace_counts))
    unexpected_trace = sorted(set(trace_counts) - expected)
    if duplicates or missing_trace or unexpected_trace:
        raise RuntimeError(
            "Item-level traceability mismatch: "
            f"duplicates={duplicates}, missing={missing_trace}, "
            f"unexpected={unexpected_trace}"
        )

    catalog_codes = set(
        re.findall(r"^\| (SRC-[A-Z]+) \|", text, flags=re.MULTILINE)
    )
    used_codes = {
        code
        for _, sources in trace_rows
        for code in re.findall(r"SRC-[A-Z]+", sources)
    }
    unknown_codes = sorted(used_codes - catalog_codes)
    unused_codes = sorted(catalog_codes - used_codes)
    if unknown_codes or unused_codes:
        raise RuntimeError(
            f"Source catalog mismatch: unknown={unknown_codes}, unused={unused_codes}"
        )

    image_links = re.findall(r"!\[[^\]]*]\(([^)]+)\)", text)
    missing_images = [
        link
        for link in image_links
        if not (SOURCE.parent / link).resolve().is_file()
    ]
    if len(image_links) != 6 or missing_images:
        raise RuntimeError(
            f"Design evidence mismatch: count={len(image_links)}, "
            f"missing={missing_images}"
        )

    external_links = sorted(set(re.findall(r"https://[^)\s]+", text)))
    if len(external_links) != 8:
        raise RuntimeError(
            f"Expected 8 official external references, found {len(external_links)}"
        )

    return {
        "requirement_ids": len(expected),
        "traceability_rows": len(trace_rows),
        "source_codes": sorted(catalog_codes),
        "design_evidence_images": len(image_links),
        "official_external_links": external_links,
    }


def add_finding(
    findings: list[dict[str, str]],
    severity: str,
    code: str,
    message: str,
) -> None:
    findings.append({"severity": severity, "code": code, "message": message})


def validate_docx_accessibility() -> tuple[dict[str, Any], dict[str, Any]]:
    findings: list[dict[str, str]] = []
    with zipfile.ZipFile(DOCX) as package:
        corrupt_member = package.testzip()
        if corrupt_member:
            add_finding(
                findings,
                "high",
                "corrupt-package",
                f"Corrupt DOCX member: {corrupt_member}",
            )

        document_xml = ElementTree.fromstring(package.read("word/document.xml"))
        doc_properties = document_xml.findall(f".//{{{WP_NS}}}docPr")
        missing_alt = [
            node.attrib.get("id", "unknown")
            for node in doc_properties
            if not (node.attrib.get("descr") or node.attrib.get("title"))
        ]
        if missing_alt:
            add_finding(
                findings,
                "high",
                "missing-alt-text",
                f"Drawing objects without alternative text: {missing_alt}",
            )

        relationships = ElementTree.fromstring(
            package.read("word/_rels/document.xml.rels")
        )
        external_targets = [
            node.attrib.get("Target", "")
            for node in relationships.findall(f"{{{REL_NS}}}Relationship")
            if node.attrib.get("TargetMode") == "External"
        ]

    document = Document(DOCX)
    if not document.core_properties.title:
        add_finding(
            findings,
            "medium",
            "missing-title",
            "Document core title is empty",
        )
    if not document.core_properties.subject:
        add_finding(
            findings,
            "medium",
            "missing-subject",
            "Document core subject is empty",
        )

    missing_table_headers: list[int] = []
    empty_table_headers: list[int] = []
    for index, table in enumerate(document.tables, start=1):
        header_property = table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader"))
        if header_property is None:
            missing_table_headers.append(index)
        if any(not cell.text.strip() for cell in table.rows[0].cells):
            empty_table_headers.append(index)
    if missing_table_headers:
        add_finding(
            findings,
            "high",
            "missing-table-header",
            f"Tables without repeated header rows: {missing_table_headers}",
        )
    if empty_table_headers:
        add_finding(
            findings,
            "medium",
            "empty-table-header",
            f"Tables with empty header cells: {empty_table_headers}",
        )

    previous_level = 0
    heading_skips: list[str] = []
    for paragraph in document.paragraphs:
        match = re.fullmatch(r"Heading ([1-9])", paragraph.style.name or "")
        if not match:
            continue
        level = int(match.group(1))
        if previous_level and level > previous_level + 1:
            heading_skips.append(paragraph.text)
        previous_level = level
    if heading_skips:
        add_finding(
            findings,
            "high",
            "heading-level-skip",
            f"Heading hierarchy skips levels at: {heading_skips}",
        )

    if len(external_targets) != 8:
        add_finding(
            findings,
            "medium",
            "external-link-count",
            f"Expected 8 official external hyperlinks, found {len(external_targets)}",
        )

    counts = Counter(finding["severity"] for finding in findings)
    accessibility_report = {
        "file": str(DOCX.relative_to(ROOT)),
        "counts": {
            "high": counts["high"],
            "medium": counts["medium"],
            "low": counts["low"],
        },
        "findings": findings,
    }
    structural_report = {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "drawing_objects_with_alt_text": len(doc_properties) - len(missing_alt),
        "external_hyperlinks": len(external_targets),
    }
    return accessibility_report, structural_report


def prepare_render_dir(path: Path) -> Path:
    resolved = path.expanduser().resolve()
    if resolved.exists() and any(resolved.iterdir()):
        raise RuntimeError(
            f"Render directory must be new or empty to avoid overwriting evidence: {resolved}"
        )
    resolved.mkdir(parents=True, exist_ok=True)
    return resolved


def render_docx(render_dir: Path) -> dict[str, Any]:
    soffice = require_command("soffice", "libreoffice")
    pdftoppm = require_command("pdftoppm")

    with tempfile.TemporaryDirectory(prefix="evo-tz-libreoffice-profile-") as profile:
        profile_uri = Path(profile).resolve().as_uri()
        conversion = run(
            [
                soffice,
                f"-env:UserInstallation={profile_uri}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(render_dir),
                str(DOCX),
            ]
        )

    pdf_path = render_dir / f"{DOCX.stem}.pdf"
    if not pdf_path.is_file():
        raise RuntimeError(
            "LibreOffice did not create the expected PDF. "
            f"stdout={conversion.stdout!r}, stderr={conversion.stderr!r}"
        )

    run(
        [
            pdftoppm,
            "-png",
            "-r",
            "150",
            str(pdf_path),
            str(render_dir / "page"),
        ]
    )
    page_paths = sorted(
        render_dir.glob("page-*.png"),
        key=lambda path: int(path.stem.rsplit("-", 1)[1]),
    )
    if not page_paths:
        raise RuntimeError("Poppler did not render any DOCX pages")

    return {
        "soffice": soffice,
        "pdftoppm": pdftoppm,
        "pdf": str(pdf_path),
        "pages": len(page_paths),
        "page_files": [path.name for path in page_paths],
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build, render and verify EVO_PLATFORM_TZ.docx"
    )
    parser.add_argument(
        "--render-dir",
        type=Path,
        required=True,
        help="New or empty directory for the real PDF/PNG validation evidence",
    )
    parser.add_argument(
        "--skip-build",
        action="store_true",
        help="Verify the existing DOCX instead of rebuilding it first",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    render_dir = prepare_render_dir(args.render_dir)

    reproducibility = {
        "checked": False,
        "bit_for_bit": None,
    }
    if not args.skip_build:
        run([sys.executable, str(GENERATOR)])
        first_hash = sha256(DOCX)
        run([sys.executable, str(GENERATOR)])
        second_hash = sha256(DOCX)
        reproducibility = {
            "checked": True,
            "bit_for_bit": first_hash == second_hash,
            "first_sha256": first_hash,
            "second_sha256": second_hash,
        }
        if first_hash != second_hash:
            raise RuntimeError(
                "Two consecutive DOCX builds produced different SHA-256 hashes"
            )
    if not DOCX.is_file():
        raise RuntimeError(f"Generated DOCX is missing: {DOCX}")

    markdown = validate_markdown()
    accessibility, structural = validate_docx_accessibility()
    render = render_docx(render_dir)

    a11y_path = render_dir / "accessibility-report.json"
    a11y_path.write_text(
        json.dumps(accessibility, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )

    validation = {
        "source": str(SOURCE.relative_to(ROOT)),
        "docx": str(DOCX.relative_to(ROOT)),
        "python": sys.version.split()[0],
        "python_docx": docx.__version__,
        "source_sha256": sha256(SOURCE),
        "docx_sha256": sha256(DOCX),
        "docx_reproducibility": reproducibility,
        "markdown": markdown,
        "docx_structure": structural,
        "accessibility": accessibility["counts"],
        "render": render,
        "human_visual_inspection": {
            "required": True,
            "status": "record separately in docs/specs/EVO_PLATFORM_TZ_VALIDATION.md",
        },
    }
    validation_path = render_dir / "validation.json"
    validation_path.write_text(
        json.dumps(validation, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )

    if any(accessibility["counts"].values()):
        print(json.dumps(accessibility, ensure_ascii=False, indent=2))
        raise RuntimeError("DOCX accessibility audit contains unresolved findings")

    print(
        "PASS "
        f"requirements={markdown['requirement_ids']} "
        f"traceability={markdown['traceability_rows']} "
        f"pages={render['pages']} "
        "a11y=0/0/0 "
        f"evidence={render_dir}"
    )
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except (OSError, RuntimeError, subprocess.CalledProcessError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        raise SystemExit(1)
